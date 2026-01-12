# Advanced Document Content Analysis and Extraction System
"""
Intelligent document processing module supporting multiple file formats
with sophisticated content extraction and structural analysis capabilities.
"""

import os
from pathlib import Path
from typing import Dict, List, Any, Optional
import pdfplumber
import docx
from processors.exceptions import DocumentParsingException, UnsupportedDocumentType

class DocumentProcessor:
    """
    Sophisticated document analysis engine capable of processing various 
    document formats and extracting structured content with metadata.
    """
    
    SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".md", ".txt"}
    
    @classmethod
    def analyze_document_content(cls, document_path: str) -> Dict[str, Any]:
        """
        Primary document analysis interface that routes to appropriate
        format-specific processors based on file extension detection.
        
        Args:
            document_path: File system path to the document for analysis
            
        Returns:
            Dictionary containing extracted text and structural metadata
            
        Raises:
            UnsupportedDocumentType: When file format is not supported
            DocumentParsingException: When content extraction fails
        """
        file_extension = Path(document_path).suffix.lower()
        
        try:
            if file_extension == ".pdf":
                return cls._process_pdf_document(document_path)
            elif file_extension == ".docx":
                return cls._process_word_document(document_path)
            elif file_extension in [".md", ".txt"]:
                return cls._process_text_document(document_path)
            else:
                raise UnsupportedDocumentType(
                    f"Document format '{file_extension}' is not currently supported. "
                    f"Supported formats: {', '.join(cls.SUPPORTED_EXTENSIONS)}"
                )
        except Exception as processing_error:
            raise DocumentParsingException(
                f"Critical error during document analysis of {document_path}: {processing_error}"
            )
    
    @classmethod
    def _process_pdf_document(cls, pdf_path: str) -> Dict[str, Any]:
        """
        Advanced PDF content extraction using intelligent text mining algorithms.
        Handles multi-page documents with layout preservation.
        """
        try:
            with pdfplumber.open(pdf_path) as pdf_document:
                # Extract text from all pages with intelligent spacing
                extracted_pages = []
                for page_number, page in enumerate(pdf_document.pages, 1):
                    page_text = page.extract_text()
                    if page_text:
                        extracted_pages.append({
                            "page_number": page_number,
                            "content": page_text.strip()
                        })
                
                # Combine all pages into unified text
                unified_text = "\n\n".join(page["content"] for page in extracted_pages if page["content"])
                
                return {
                    "extracted_text": unified_text,
                    "document_sections": cls._identify_content_sections(unified_text),
                    "metadata": {
                        "document_type": "pdf",
                        "total_pages": len(pdf_document.pages),
                        "processed_pages": len(extracted_pages)
                    }
                }
        except Exception as pdf_error:
            raise DocumentParsingException(f"PDF processing encountered error: {pdf_error}")
    
    @classmethod
    def _process_word_document(cls, docx_path: str) -> Dict[str, Any]:
        """
        Microsoft Word document processing with paragraph structure preservation.
        Extracts both content and formatting metadata.
        """
        try:
            word_document = docx.Document(docx_path)
            
            # Extract paragraph content with structure preservation
            paragraph_contents = []
            for paragraph in word_document.paragraphs:
                if paragraph.text.strip():
                    paragraph_contents.append(paragraph.text.strip())
            
            unified_content = "\n".join(paragraph_contents)
            
            return {
                "extracted_text": unified_content,
                "document_sections": cls._identify_content_sections(unified_content),
                "metadata": {
                    "document_type": "word_document",
                    "paragraph_count": len(paragraph_contents),
                    "total_paragraphs": len(word_document.paragraphs)
                }
            }
        except Exception as docx_error:
            raise DocumentParsingException(f"Word document processing failed: {docx_error}")
    
    @classmethod
    def _process_text_document(cls, text_path: str) -> Dict[str, Any]:
        """
        Plain text and markdown document processing with encoding detection.
        Supports multiple text encodings with fallback mechanisms.
        """
        encodings_to_try = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings_to_try:
            try:
                with open(text_path, "r", encoding=encoding) as text_file:
                    document_content = text_file.read()
                
                return {
                    "extracted_text": document_content,
                    "document_sections": cls._identify_content_sections(document_content),
                    "metadata": {
                        "document_type": "text_document",
                        "encoding_used": encoding,
                        "character_count": len(document_content)
                    }
                }
            except UnicodeDecodeError:
                continue
        
        raise DocumentParsingException(f"Unable to decode text document with any supported encoding")
    
    @classmethod
    def _identify_content_sections(cls, document_text: str) -> List[str]:
        """
        Intelligent content sectioning using header detection algorithms.
        Identifies logical document structure for enhanced processing.
        """
        import re
        
        # Advanced header detection patterns
        header_patterns = [
            r'\n#{1,6}\s+(.+)',  # Markdown headers
            r'\n([A-Z][A-Z\s]{3,})\n',  # ALL CAPS headers
            r'\n(\d+\.?\s+[A-Z][^.\n]{10,})\n',  # Numbered sections
        ]
        
        sections = []
        for pattern in header_patterns:
            matches = re.findall(pattern, document_text, re.MULTILINE)
            sections.extend(matches)
        
        # Split text by identified sections if available
        if not sections:
            # Fallback: split by double newlines for basic sectioning
            potential_sections = re.split(r'\n\s*\n', document_text)
            sections = [section.strip() for section in potential_sections if section.strip()]
        
        return sections[:10]  # Limit to first 10 sections for performance
